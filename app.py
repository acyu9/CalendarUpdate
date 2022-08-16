from docx import Document

doc = Document("AP2021.docx")
table = doc.tables[0]

# Each row has 3 columns
# Date format is Mon Aug. 17-18, 2021


# All text from column 1, cell 2
#print(doc.tables[0].rows[0].cells[1].text)

# Can't do 1 or 2 b/c 11 would be 01
int_dates = []
ones = []
for i in range(10, 32):
    int_dates.append(i)

for i in range(2, 10):
    ones.append(i)

dates = [str(x) for x in int_dates]
ones_dates = [str(x) for x in ones]

for row in table.rows:
    for cell in row.cells:
        for date in dates:
            if date in cell.paragraphs[0].text:
                new_date = str(int(date) - 1)
                run = cell.paragraphs[0].runs
                for i in range(len(run)):
                    if date in run[i].text:
                        run[i].text = run[i].text.replace(date, new_date)

            else:
                for one in ones_dates:
                    if one in cell.paragraphs[0].text:
                        new_one = str(int(one) - 1)
                        run = cell.paragraphs[0].runs
                        for i in range(len(run)):
                            if one in run[i].text:
                                run[i].text = run[i].text.replace(one, new_one)
                            elif '1' in run[i].text:
                                run[i].text = run[i].text.replace('1', '31')
                            break;
                    break;
                    
        if '1920' or '1010' in cell.text:
            print("one")
            run = cell.paragraphs[0].runs
            for i in range(len(run)):
                if '1920' or '1010' in run[i].text:
                    print("two")
                    run[i].text = run[i].text.replace('1920' or '1010', '2022')


        # prints everything in cell
        #print(cell.paragraphs[0].text)

doc.save("AP2022.docx")